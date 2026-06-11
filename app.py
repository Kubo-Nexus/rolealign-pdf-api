import os
import re
import base64
import urllib.request
from difflib import SequenceMatcher
from io import BytesIO

from flask import Flask, request, send_file, jsonify
from flask_cors import CORS

from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, white
from reportlab.lib.utils import ImageReader, simpleSplit
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = Flask(__name__)
CORS(app)

W, H = A4
API_VERSION = "1.4.0-impact-sidebar-paginate"

ACRONYMS = {
    "sap": "SAP",
    "sd": "SD",
    "mm": "MM",
    "wh": "WH",
    "fi": "FI",
    "sop": "SOP",
    "sops": "SOPs",
    "kpi": "KPI",
    "kpis": "KPIs",
    "ats": "ATS",
    "ai": "AI",
    "cv": "CV",
    "api": "API",
    "apis": "APIs",
}

SAP_SUPER_USER = "SAP Super User (SD, MM, WH & FI)"
SAP_MARKER = "__ROLEALIGN_SAP_SUPER_USER__"


def safe_hex(colour_str, fallback):
    try:
        if colour_str and isinstance(colour_str, str) and colour_str.startswith("#") and len(colour_str) in (4, 7):
            return HexColor(colour_str)
    except Exception:
        pass
    return HexColor(fallback)


def clean_text(value):
    if value is None:
        return ""

    text = str(value).strip()
    text = text.replace("\u2013", "–").replace("\u2014", "—")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+-\s*$", "", text)
    text = re.sub(r"\s+\(\s*\)\s*$", "", text)

    # Preserve the SAP skill exactly. This prevents comma splitting/normalisation damage.
    text = re.sub(
        r"\bSAP\s+Super\s+User\s*\(\s*SD\s*,?\s*MM\s*,?\s*WH\s*&\s*FI\s*\)",
        SAP_SUPER_USER,
        text,
        flags=re.IGNORECASE,
    )
    text = text.replace("Sd Mm Wh & Fi", "SD, MM, WH & FI")
    text = text.replace("SD MM WH & FI", "SD, MM, WH & FI")

    # Do not let the acronym pass damage already-correct SAP Super User text.
    protected = text.replace(SAP_SUPER_USER, SAP_MARKER)
    for wrong, right in ACRONYMS.items():
        protected = re.sub(rf"\b{wrong}\b", right, protected, flags=re.IGNORECASE)
    text = protected.replace(SAP_MARKER, SAP_SUPER_USER)

    return text.strip()


def clean_join(parts, sep=" - "):
    cleaned = [clean_text(p) for p in parts if clean_text(p)]
    return sep.join(cleaned)


def repair_skill_fragments(raw_skills):
    """Repair known skill-fragment cases before rendering.

    Some extraction/payload paths can split parenthetical acronym groups into
    separate list items, e.g. ["SAP Super User (SD", "MM", "WH & FI)"].
    Paid templates must render this as one customer-facing skill.
    """
    parts = []
    for item in raw_skills or []:
        if isinstance(item, dict):
            item = item.get("name") or item.get("skill") or item.get("label")
        txt = clean_text(item)
        if txt:
            parts.append(txt)

    repaired = []
    i = 0
    while i < len(parts):
        current = parts[i]
        lower = current.lower()

        if "sap super user" in lower and "(" in current and ")" not in current:
            group = [current]
            j = i + 1
            while j < len(parts) and len(group) < 6:
                group.append(parts[j])
                if ")" in parts[j]:
                    break
                j += 1

            combined = ", ".join(group)
            combined = re.sub(r"\s*,\s*", ", ", combined)
            combined = re.sub(r"\(\s*SD\s*,\s*MM\s*,\s*WH\s*,\s*FI\s*\)", "(SD, MM, WH & FI)", combined, flags=re.IGNORECASE)
            combined = re.sub(r"\(\s*SD\s*,\s*MM\s*,\s*WH\s*&\s*FI\s*\)", "(SD, MM, WH & FI)", combined, flags=re.IGNORECASE)
            combined = re.sub(r"\bSAP\s+Super\s+User\s*\(\s*SD\s*,\s*MM\s*,\s*WH\s*&\s*FI\s*\)", SAP_SUPER_USER, combined, flags=re.IGNORECASE)
            repaired.append(clean_text(combined))
            i = j + 1
            continue

        repaired.append(current)
        i += 1

    return repaired


def _normalise_for_compare(text):
    """Aggressive normalisation used only for duplicate detection."""
    if not text:
        return ""
    t = str(text).lower().strip()
    t = re.sub(r"[^\w\s%]", " ", t)
    t = re.sub(r"\s+", " ", t)
    # Common paraphrase equivalences
    t = t.replace("r220 million", "r220m").replace("r140 million", "r140m")
    t = t.replace("r8 million", "r8m")
    t = re.sub(r"\b(the|a|an|of|in|to|for|on|at|by|with|from|across|through|including|approximately|annually|while|that|which|and)\b", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


_STOPLIKE_TOKENS = {"all", "office", "company", "consistently", "year", "years", "first", "second", "one", "two", "performance", "successfully", "improvements", "improvement", "within", "online", "as", "saving", "savings", "generating", "manager", "marketing", "commercial", "revenue", "while", "during", "over", "monthly", "won", "awarded", "received", "achieved"}


def _significant_tokens(norm_text):
    return {tok for tok in norm_text.split() if tok and tok not in _STOPLIKE_TOKENS and not tok.isdigit()}


def dedup_items(items, similarity_threshold=0.78):
    """De-duplicate a list of strings using multi-strategy matching.

    Conservative thresholds: catches obvious paraphrases (same numeric values
    + same key tokens) but preserves distinct items that share only generic
    framing words (e.g. "Renault Aftersales World Conference" vs "DHL World
    Conference" — both panel-member events but at different organisations).
    """
    if not items:
        return []

    cleaned = []
    for raw in items:
        text = clean_text(raw) if not isinstance(raw, str) else raw.strip()
        if text:
            cleaned.append(text)

    if not cleaned:
        return []

    kept = []
    for item in cleaned:
        norm = _normalise_for_compare(item)
        if not norm:
            continue
        tokens = _significant_tokens(norm)

        is_duplicate = False
        for idx, (existing_text, existing_norm, existing_tokens) in enumerate(kept):
            duplicate_match = False

            if norm == existing_norm:
                duplicate_match = True
            elif norm in existing_norm or existing_norm in norm:
                duplicate_match = True
            elif tokens and existing_tokens:
                smaller, larger = (tokens, existing_tokens) if len(tokens) <= len(existing_tokens) else (existing_tokens, tokens)
                if len(smaller) >= 3:
                    overlap = len(smaller & larger) / len(smaller)
                    if overlap >= 0.80:
                        duplicate_match = True
                if not duplicate_match:
                    ratio = SequenceMatcher(None, norm, existing_norm).ratio()
                    if ratio >= similarity_threshold:
                        duplicate_match = True

            if duplicate_match:
                is_duplicate = True
                if len(item) > len(existing_text):
                    kept[idx] = (item, norm, tokens)
                break

        if not is_duplicate:
            kept.append((item, norm, tokens))

    return [text for text, _, _ in kept]


def normalise_references(value):
    """Clean references defensively.

    Base44 can sometimes send references as a list of single characters after
    sanitisation, which produced output like "A; v; a; i...". Render must
    never show that. Keep normal strings/lists, but collapse character lists
    and canonicalise common reference phrases.
    """
    if value is None:
        return ""
    if isinstance(value, list):
        vals = [clean_text(v) for v in value if clean_text(v)]
        if vals and sum(1 for v in vals if len(v) <= 2) >= max(4, int(len(vals) * 0.6)):
            text = "".join(vals)
        else:
            text = "; ".join(vals)
    else:
        text = clean_text(value)
    text = re.sub(r"\s*;\s*([A-Za-z])\s*;\s*", r"\1", text)
    text = re.sub(r"\bA\s*v\s*a\s*i\s*l\s*a\s*b\s*l\s*e\b", "Available", text, flags=re.IGNORECASE)
    text = re.sub(r"\bu\s*p\s*o\s*n\b", "upon", text, flags=re.IGNORECASE)
    text = re.sub(r"\br\s*e\s*q\s*u\s*e\s*s\s*t\b", "request", text, flags=re.IGNORECASE)
    text = clean_text(text)
    if not text:
        return ""
    # Privacy / POPIA: never echo referee names, phone numbers, or email
    # addresses into a rendered CV. Any non-empty references value is collapsed
    # to a single neutral line. Enforced here so every template and the DOCX
    # inherit the behaviour from one place.
    return "Available upon request."


def split_skills(skills):
    if isinstance(skills, str):
        protected = re.sub(
            r"\bSAP\s+Super\s+User\s*\(\s*SD\s*,?\s*MM\s*,?\s*WH\s*&\s*FI\s*\)",
            SAP_MARKER,
            skills,
            flags=re.IGNORECASE,
        )
        # Split common skill separators, but avoid breaking the SAP parenthetical commas.
        raw_parts = re.split(r"\n|•|·|;|,(?!\s*(?:MM|WH|FI)\b)", protected)
        return [p.replace(SAP_MARKER, SAP_SUPER_USER) for p in raw_parts]
    return repair_skill_fragments(skills or [])


def normalise_cv(cv):
    cv = cv or {}
    name = clean_text(cv.get("name") or cv.get("full_name") or "")

    clean_cv = {
        "name": name,
        "email": clean_text(cv.get("email")),
        "phone": clean_text(cv.get("phone")),
        "location": clean_text(cv.get("location")),
        "linkedin": clean_text(cv.get("linkedin")),
        "summary": clean_text(cv.get("summary")),
        "photo": cv.get("photo") or cv.get("profile_photo") or cv.get("photo_url") or cv.get("image"),
        "is_premium": bool(cv.get("is_premium", True)),
        "render_version": cv.get("render_version") or API_VERSION,
    }

    skills = split_skills(cv.get("skills") or [])
    clean_skills = []
    for skill in skills:
        if isinstance(skill, dict):
            skill = skill.get("name") or skill.get("skill") or skill.get("label")
        s = clean_text(skill)
        if s and s not in clean_skills:
            clean_skills.append(s)
    clean_cv["skills"] = clean_skills
    clean_cv["skills_csv"] = ", ".join(clean_skills)
    clean_cv["skills_bullet"] = " · ".join(clean_skills)

    experience = []
    for job in cv.get("experience", []) or []:
        if isinstance(job, str):
            title = clean_text(job)
            company = ""
            dates = ""
            bullets = []
        else:
            title = clean_text(job.get("title") or job.get("role") or job.get("position"))
            company = clean_text(job.get("company") or job.get("employer") or job.get("organisation") or job.get("organization"))
            dates = clean_text(job.get("dates") or job.get("date") or job.get("period"))
            bullets = [clean_text(b) for b in job.get("bullets", []) if clean_text(b)]

        header = clean_text(job.get("header")) if isinstance(job, dict) else ""
        if not header:
            left = clean_join([title, company], " — ")
            header = clean_join([left, dates], " · ")

        if title or company or dates or bullets:
            experience.append(
                {
                    "title": title,
                    "company": company,
                    "dates": dates,
                    "header": header,
                    "bullets": bullets,
                }
            )
    clean_cv["experience"] = experience

    education = []
    for edu in cv.get("education", []) or []:
        if isinstance(edu, str):
            degree = clean_text(edu)
            institution = ""
            year = ""
        else:
            degree = clean_text(edu.get("degree") or edu.get("qualification") or edu.get("name"))
            institution = clean_text(edu.get("institution") or edu.get("school") or edu.get("provider"))
            year = clean_text(edu.get("year") or edu.get("dates") or edu.get("date"))

        line = clean_join([degree, institution], " - ")
        if year:
            line = f"{line} ({year})" if line else year

        if line and line not in ["Professional Certifications", "Professional Certification"]:
            education.append({"degree": degree, "institution": institution, "year": year, "line": line})
        elif degree and "professional certification" not in degree.lower():
            education.append({"degree": degree, "institution": institution, "year": year, "line": degree})
    clean_cv["education"] = education

    certifications = []
    for cert in cv.get("certifications", []) or []:
        cert_text = clean_text(cert.get("name") if isinstance(cert, dict) else cert)
        if cert_text and cert_text not in ["Professional Certifications", "Professional Certification"]:
            certifications.append(cert_text)
    clean_cv["certifications"] = dedup_items(certifications)

    achievements_raw = (
        cv.get("achievements")
        or cv.get("key_achievements")
        or cv.get("career_achievements")
        or cv.get("accomplishments")
        or []
    )
    if isinstance(achievements_raw, str):
        achievements_raw = re.split(r"\n|•|;", achievements_raw)
    raw_achievements_list = [clean_text(a.get("text") if isinstance(a, dict) else a) for a in achievements_raw if clean_text(a.get("text") if isinstance(a, dict) else a)]
    # Also merge any per-role key_achievements from the experience array
    for exp_entry in cv.get("experience", []) or []:
        if isinstance(exp_entry, dict):
            for role_ach in (exp_entry.get("key_achievements") or exp_entry.get("achievements") or []):
                ach_text = clean_text(role_ach.get("text") if isinstance(role_ach, dict) else role_ach)
                if ach_text:
                    raw_achievements_list.append(ach_text)
    clean_cv["achievements"] = dedup_items(raw_achievements_list)

    systems_raw = (
        cv.get("systems")
        or cv.get("systems_experience")
        or cv.get("technical_skills")
        or cv.get("tools")
        or []
    )
    if isinstance(systems_raw, str):
        systems_raw = re.split(r"\n|•|;|,(?!\s*(?:MM|WH|FI)\b)", systems_raw)
    clean_cv["systems_experience"] = dedup_items(repair_skill_fragments([clean_text(s.get("name") if isinstance(s, dict) else s) for s in systems_raw if clean_text(s.get("name") if isinstance(s, dict) else s)]))

    references_raw = cv.get("references") or cv.get("reference") or ""
    clean_cv["references"] = normalise_references(references_raw)

    return clean_cv


def decode_photo(photo_data):
    if not photo_data:
        return None
    try:
        if isinstance(photo_data, str) and photo_data.startswith("http"):
            req = urllib.request.Request(photo_data, headers={"User-Agent": f"RoleAlignPDF/{API_VERSION}"})
            with urllib.request.urlopen(req, timeout=8) as response:
                return ImageReader(BytesIO(response.read()))

        data = str(photo_data)
        if "," in data:
            data = data.split(",", 1)[1]
        img_bytes = base64.b64decode(data)
        return ImageReader(BytesIO(img_bytes))
    except Exception as e:
        print(f"[photo_error] {e}")
        return None


def draw_circular_photo(c, img_reader, cx, cy, radius):
    if img_reader is None:
        return
    c.saveState()
    path = c.beginPath()
    path.circle(cx, cy, radius)
    path.close()
    c.clipPath(path, stroke=0)
    c.drawImage(
        img_reader,
        cx - radius,
        cy - radius,
        radius * 2,
        radius * 2,
        preserveAspectRatio=True,
        mask="auto",
    )
    c.restoreState()


def draw_rolealign_watermark(c):
    c.saveState()
    try:
        c.setFillColor(HexColor("#D9D9D9"))
        c.setFillAlpha(0.18)
    except Exception:
        c.setFillColor(HexColor("#E6E6E6"))
    c.setFont("Helvetica-Bold", 58)
    c.translate(W / 2, H / 2)
    c.rotate(35)
    c.drawCentredString(0, 0, "RoleAlign")
    c.drawCentredString(0, -110, "RoleAlign")
    c.drawCentredString(0, 110, "RoleAlign")
    c.restoreState()


def footer_brand(c, is_premium):
    """Draw branding only on free/watermarked outputs.

    Paid PDFs must look like normal customer CVs, so they get no footer
    and no visible RoleAlign mark. Free Starter keeps both watermark and
    small attribution footer.
    """
    if is_premium:
        return
    draw_rolealign_watermark(c)
    c.setFont("Helvetica", 6)
    c.setFillColor(HexColor("#999999"))
    c.drawCentredString(W / 2, 12, "Created with RoleAlign")


def section_heading(c, x, y, label, colour, width=80):
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(colour)
    c.drawString(x, y, label)
    c.setStrokeColor(colour)
    c.setLineWidth(0.8)
    c.line(x, y - 4, x + width, y - 4)


def wrap_lines(text, font_name, font_size, width):
    """Wrap English text safely by words using ReportLab's width metrics.

    This avoids ReportLab paragraph wrapping modes, which caused visible internal
    spacing/splitting in generated PDFs. Only extremely long single tokens
    are allowed to break as a last resort by simpleSplit.
    """
    text = clean_text(text)
    if not text:
        return []
    return simpleSplit(text, font_name, font_size, width)


def draw_wrapped(
    c,
    text,
    x,
    y,
    width,
    font="Helvetica",
    size=8,
    leading=11,
    colour="#4A4A4A",
    bold=False,
):
    text = clean_text(text)
    if not text:
        return y

    font_name = "Helvetica-Bold" if bold else font
    lines = wrap_lines(text, font_name, size, width)
    c.setFont(font_name, size)
    c.setFillColor(HexColor(colour))

    baseline = y - size
    for line in lines:
        c.drawString(x, baseline, line)
        baseline -= leading

    return y - (leading * len(lines))


def draw_manual_bullet(c, text, x, y, width, colour="#4A4A4A", size=8, leading=11, bullet=True):
    text = clean_text(text)
    if not text:
        return y

    c.setFont("Helvetica", size)
    c.setFillColor(HexColor(colour))

    if bullet:
        bullet_x = x
        text_x = x + 10
        text_width = width - 15
        lines = wrap_lines(text, "Helvetica", size, text_width)
        if not lines:
            return y

        first_baseline = y - size
        c.drawString(bullet_x, first_baseline, "•")
        c.drawString(text_x, first_baseline, lines[0])

        baseline = first_baseline - leading
        for line in lines[1:]:
            c.drawString(text_x, baseline, line)
            baseline -= leading

        return y - (leading * len(lines))

    lines = wrap_lines(text, "Helvetica", size, width)
    baseline = y - size
    for line in lines:
        c.drawString(x, baseline, line)
        baseline -= leading
    return y - (leading * len(lines))


def draw_role(
    c,
    job,
    x,
    y,
    width,
    accent,
    text_dark="#1A1A1A",
    text_med="#4A4A4A",
    text_light="#777777",
    bullet=True,
    company_gap=9,
    new_page=None,
    bullet_bottom=70,
):
    """
    Stable role header renderer.

    The company baseline is calculated from the actual wrapped title paragraph height,
    not a fixed one-line offset. This prevents wrapped titles from colliding
    with company names.

    If a `new_page` callback is supplied, the role paginates instead of
    clipping: it breaks before drawing the header if the title would be
    orphaned at the bottom of the page, and before any bullet that would
    otherwise run off the page. `new_page` must start a fresh page (drawing
    the template's continuation shell) and return the new starting y.
    """
    title = clean_text(job.get("title"))
    company = clean_text(job.get("company"))
    dates = clean_text(job.get("dates"))

    # Avoid orphaning a role header at the very bottom of a page.
    if new_page is not None and y < 118:
        y = new_page()

    title_w = width - 92 if dates else width
    title_top_y = y
    title_bottom_y = y

    if title:
        title_bottom_y = draw_wrapped(
            c,
            title,
            x,
            title_top_y,
            title_w,
            size=8.8,
            leading=11.2,
            colour=text_dark,
            bold=True,
        )

    if dates:
        c.setFont("Helvetica", 7.5)
        c.setFillColor(HexColor(text_light))
        tw = c.stringWidth(dates, "Helvetica", 7.5)
        c.drawString(x + width - tw, title_top_y - 8.5, dates)

    # Company must sit below the full wrapped title block, never inline with title line 2.
    y = min(title_bottom_y - company_gap, title_top_y - 16)

    if company:
        c.setFont("Helvetica-Bold", 8)
        c.setFillColor(accent)
        c.drawString(x, y, company)
        y -= 16
    else:
        y -= 3

    for item in job.get("bullets", []) or []:
        txt = clean_text(item)
        if not txt:
            continue
        if new_page is not None and y < bullet_bottom:
            y = new_page()
        y = draw_manual_bullet(
            c,
            txt,
            x,
            y,
            width - 5,
            colour=text_med,
            size=8,
            leading=11,
            bullet=bullet,
        )
        y -= 3


    return y - 8


def draw_skills_list(c, skills, x, y, width, colour="#4A4A4A", size=7.5):
    for skill in skills:
        text = clean_text(skill)
        if not text:
            continue
        y = draw_wrapped(
            c,
            text,
            x,
            y,
            width,
            font="Helvetica",
            size=size,
            leading=size + 2,
            colour=colour,
            bold=False,
        )
        y -= 5
    return y


def draw_skill_pills(c, skills, x, y, width, bg, fg, font_size=7, min_y=None):
    # One pill per skill, anchored from the TOP and grown downward, with a
    # fixed gap between pills. This guarantees a multi-line pill (e.g. a long
    # skill that wraps) can never overlap and merge into the pill above it.
    #
    # When min_y is supplied, the column is height-aware: it stops before any
    # pill would cross the floor and returns the skills it could not fit, so the
    # caller can continue them in the next page's sidebar instead of letting
    # them run off the bottom of the page. Returns (end_y, remaining_skills).
    line_height = font_size + 3
    v_pad = 5
    gap = 6

    skills = [s for s in (skills or []) if clean_text(s)]
    remaining = []
    for idx, skill in enumerate(skills):
        skill_clean = clean_text(skill)

        lines = wrap_lines(skill_clean, "Helvetica", font_size, width - 14)
        if not lines:
            continue
        pill_h = len(lines) * line_height + 2 * v_pad

        # Stop before drawing a pill that would cross the bottom floor; hand the
        # rest back so the caller can continue them on the next page's sidebar.
        if min_y is not None and (y - pill_h) < min_y:
            remaining = skills[idx:]
            break

        top = y
        c.setFillColor(bg)
        c.roundRect(x, top - pill_h, width, pill_h, 8, fill=1, stroke=0)
        c.setFillColor(fg)
        c.setFont("Helvetica", font_size)
        baseline = top - v_pad - font_size + 1
        for line in lines:
            c.drawString(x + 7, baseline, line)
            baseline -= line_height
        y = top - pill_h - gap

    return y - 6, remaining


def draw_education(c, education, x, y, width, title_colour="#1A1A1A", sub_colour="#777777"):
    for edu in education:
        line = clean_text(edu.get("line"))
        if not line:
            continue
        y = draw_wrapped(c, line, x, y, width, size=7.2, leading=9, colour=title_colour, bold=False)
        y -= 7
    return y


@app.route("/", methods=["GET"])
def home():
    return jsonify(
        {
            "status": "RoleAlign PDF API is running",
            "service": "rolealign-pdf-api",
            "version": API_VERSION,
            "endpoints": {
                "GET /health": "Health check",
                "POST /generate-pdf": "Generate a styled CV PDF",
                "POST /generate-docx": "Generate an editable CV DOCX",
            },
        }
    )


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"ok": True, "service": "rolealign-pdf-api", "version": API_VERSION})


def draw_list_section(c, label, items, x, y, width, accent, text_colour="#4A4A4A", bullet=True):
    items = [clean_text(i) for i in (items or []) if clean_text(i)]
    if not items:
        return y
    section_heading(c, x, y, label, accent, min(width, 150))
    y -= 18
    for item in items:
        y = draw_manual_bullet(
            c,
            item,
            x,
            y,
            width,
            colour=text_colour,
            size=8,
            leading=11,
            bullet=bullet,
        )
        y -= 3
    return y - 10


def generate_starter_pdf(cv, colours):
    cv = normalise_cv(cv)
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)

    x = 42
    y = H - 45
    width = W - 84
    accent = safe_hex(colours.get("accent"), "#111827")

    footer_brand(c, cv.get("is_premium", False))

    c.setFont("Helvetica-Bold", 22)
    c.setFillColor(HexColor("#111827"))
    c.drawString(x, y, cv.get("name") or "Professional CV")
    y -= 16

    c.setFont("Helvetica", 8)
    c.setFillColor(HexColor("#4B5563"))
    c.drawString(x, y, clean_join([cv.get("email"), cv.get("phone"), cv.get("location")], " | "))
    y -= 28

    section_heading(c, x, y, "SUMMARY", accent, 65)
    y -= 16
    y = draw_wrapped(c, cv.get("summary"), x, y, width, size=8.5, leading=12, colour="#374151") - 14

    section_heading(c, x, y, "EXPERIENCE", accent, 80)
    y -= 18
    def _exp_newpage():
        c.showPage()
        footer_brand(c, cv.get("is_premium", False))
        return H - 45
    for job in cv.get("experience", []):
        y = draw_role(c, job, x, y, width, accent, bullet=False, company_gap=9, new_page=_exp_newpage)

    if y < 110:
        c.showPage()
        footer_brand(c, cv.get("is_premium", False))
        y = H - 45

    section_heading(c, x, y, "SKILLS", accent, 45)
    y -= 16
    y = draw_wrapped(c, cv.get("skills_csv"), x, y, width, size=8, leading=11, colour="#374151") - 14

    section_heading(c, x, y, "EDUCATION", accent, 70)
    y -= 16
    y = draw_education(c, cv.get("education", []), x, y, width)

    def starter_extra(label, items):
        nonlocal y
        items = [clean_text(i) for i in (items or []) if clean_text(i)]
        if not items:
            return
        if y < 90:
            c.showPage()
            footer_brand(c, cv.get("is_premium", False))
            y = H - 45
        y -= 14
        section_heading(c, x, y, label, accent, min(width, 150))
        y -= 16
        for item in items:
            y = draw_manual_bullet(c, item, x, y, width, colour="#374151", size=8, leading=10, bullet=True)
            y -= 1
        y -= 6

    starter_extra("KEY ACHIEVEMENTS", cv.get("achievements", []))
    starter_extra("PROFESSIONAL CERTIFICATIONS", cv.get("certifications", []))
    starter_extra("SYSTEMS EXPERIENCE", cv.get("systems_experience", []))

    if cv.get("references"):
        y -= 14
        if y >= 60:
            section_heading(c, x, y, "REFERENCES", accent, 70)
            y -= 16
            draw_wrapped(c, cv.get("references"), x, y, width, size=8, leading=11, colour="#374151")

    c.save()
    buf.seek(0)
    return buf


def generate_executive_pdf(cv, colours):
    cv = normalise_cv(cv)
    SIDEBAR_W = 190
    NAVY = safe_hex(colours.get("primary"), "#1B2A4A")
    ACCENT = safe_hex(colours.get("accent"), "#C9A96E")
    TEXT_DARK = "#1A1A1A"
    TEXT_MED = "#4A4A4A"
    TEXT_LIGHT = "#7A7A7A"

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    photo_img = decode_photo(cv.get("photo"))

    def executive_page1_sidebar():
        """Render Executive sidebar only on page 1."""
        c.setFillColor(NAVY)
        c.rect(0, 0, SIDEBAR_W, H, fill=1, stroke=0)
        footer_brand(c, cv.get("is_premium", True))

        sx = 18
        cx, cy = SIDEBAR_W / 2, H - 70

        if photo_img:
            c.setFillColor(white)
            c.circle(cx, cy, 38, fill=1, stroke=0)
            draw_circular_photo(c, photo_img, cx, cy, 36)
        else:
            c.setFillColor(HexColor("#2A3F6A"))
            c.circle(cx, cy, 36, fill=1, stroke=0)
            c.setFillColor(white)
            c.setFont("Helvetica-Bold", 18)
            initials = "".join([w[0] for w in (cv.get("name") or "CV").split()[:2]]).upper()
            c.drawCentredString(cx, cy - 6, initials)

        y = H - 130

        section_heading(c, sx, y, "CONTACT", ACCENT, SIDEBAR_W - 36)
        y -= 20
        y = draw_skills_list(
            c,
            [cv.get("email"), cv.get("phone"), cv.get("location"), cv.get("linkedin")],
            sx,
            y,
            SIDEBAR_W - 36,
            "#D0D0D0",
            7.3,
        )
        y -= 10

        section_heading(c, sx, y, "SKILLS", ACCENT, SIDEBAR_W - 36)
        y -= 20
        skill_size = 6.8 if len(cv.get("skills", [])) > 14 else 7.2
        y = draw_skills_list(c, cv.get("skills", []), sx, y, SIDEBAR_W - 36, "#D0D0D0", skill_size)
        y -= 10

        if y > 95:
            section_heading(c, sx, y, "EDUCATION", ACCENT, SIDEBAR_W - 36)
            y -= 20
            draw_education(c, cv.get("education", []), sx, y, SIDEBAR_W - 36, "#FFFFFF", "#A0A0A0")

    def executive_continuation_shell(page_no):
        """Page 2+ keeps the Executive rail but no repeated sidebar sections."""
        footer_brand(c, cv.get("is_premium", True))
        c.setFillColor(NAVY)
        c.rect(0, 0, SIDEBAR_W, H, fill=1, stroke=0)
        c.setFont("Helvetica-Bold", 9)
        c.setFillColor(NAVY)
        c.drawString(SIDEBAR_W + 24, H - 24, cv.get("name") or "Professional CV")
        c.setFont("Helvetica", 7)
        c.setFillColor(HexColor(TEXT_LIGHT))
        c.drawRightString(W - 24, H - 24, f"Page {page_no}")
        c.setStrokeColor(ACCENT)
        c.setLineWidth(0.5)
        c.line(SIDEBAR_W + 24, H - 32, W - 24, H - 32)

    def draw_compact_section(label, items, x, y, width):
        """Compact Executive section that paginates mid-list rather than
        clipping. A long section flows onto a new page instead of running off
        the bottom. A short trailing page is acceptable; losing content is not.
        """
        nonlocal page_no
        items = [clean_text(i) for i in (items or []) if clean_text(i)]
        if not items:
            return y
        BOTTOM = 70  # never draw a bullet below this y; leave room for wrapping
        # Need room for the heading plus at least the first item.
        if y < BOTTOM + 34:
            c.showPage()
            page_no += 1
            executive_continuation_shell(page_no)
            y = H - 48
        y -= 14
        section_heading(c, x, y, label, NAVY, min(width, 150))
        y -= 16
        for item in items:
            if y < BOTTOM:
                c.showPage()
                page_no += 1
                executive_continuation_shell(page_no)
                y = H - 48
            y = draw_manual_bullet(c, item, x, y, width, colour=TEXT_MED, size=7.0, leading=8.0, bullet=True)
            y -= 0.8
        return y - 6

    executive_page1_sidebar()

    mx = SIDEBAR_W + 24
    mw = W - mx - 24
    y = H - 45
    page_no = 1

    c.setFont("Helvetica-Bold", 24)
    c.setFillColor(NAVY)
    c.drawString(mx, y, cv.get("name") or "Professional CV")
    y -= 26

    section_heading(c, mx, y, "SUMMARY", NAVY, mw)
    y -= 16
    y = draw_wrapped(c, cv.get("summary"), mx, y, mw, size=8.5, leading=13, colour=TEXT_MED) - 16

    section_heading(c, mx, y, "EXPERIENCE", NAVY, mw)
    y -= 18

    def _exp_newpage():
        nonlocal page_no, mx, mw
        c.showPage()
        page_no += 1
        executive_continuation_shell(page_no)
        mx = SIDEBAR_W + 24
        mw = W - mx - 24
        return H - 48
    for job in cv.get("experience", []):
        y = draw_role(c, job, mx, y, mw, ACCENT, TEXT_DARK, TEXT_MED, TEXT_LIGHT, company_gap=9, new_page=_exp_newpage)

    extra_sections = [
        ("KEY ACHIEVEMENTS", cv.get("achievements", [])),
        ("PROFESSIONAL CERTIFICATIONS", cv.get("certifications", [])),
        ("SYSTEMS EXPERIENCE", cv.get("systems_experience", [])),
    ]
    for label, items in extra_sections:
        if items:
            y = draw_compact_section(label, items, mx, y, mw)

    if cv.get("references"):
        # References are a single neutral line. Never spawn a near-empty extra
        # page just to show it — if it does not fit, omit it so the CV ends
        # cleanly on the current page.
        y -= 14  # breathing room above REFERENCES heading
        if y >= 42:
            section_heading(c, mx, y, "REFERENCES", NAVY, 75)
            y -= 16
            draw_wrapped(c, cv.get("references"), mx, y, mw, size=7.5, leading=9, colour=TEXT_MED)

    c.save()
    buf.seek(0)
    return buf

def generate_creative_pdf(cv, colours):
    cv = normalise_cv(cv)
    P1 = safe_hex(colours.get("primary_1", colours.get("primary")), "#6366F1")
    P2 = safe_hex(colours.get("primary_2", colours.get("accent")), "#8B5CF6")
    TEXT_DARK = "#1A1A2E"
    TEXT_MED = "#4A4A5A"
    RIGHT_W = 185
    LEFT_W = W - RIGHT_W - 48
    PANEL_BG = HexColor("#F0EDFF")

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    photo_img = decode_photo(cv.get("photo"))

    def draw_page1_shell():
        footer_brand(c, cv.get("is_premium", True))
        band_h = 90
        c.setFillColor(P1)
        c.rect(0, H - band_h, W, band_h, fill=1, stroke=0)

        c.setFont("Helvetica-Bold", 28)
        c.setFillColor(white)
        c.drawString(28, H - 50, cv.get("name") or "Professional CV")

        c.setFont("Helvetica", 7.5)
        c.setFillColor(HexColor("#D0CDFF"))
        c.drawString(28, H - 83, clean_join([cv.get("email"), cv.get("phone"), cv.get("location")], " | "))

        panel_x = W - RIGHT_W
        c.setFillColor(PANEL_BG)
        c.rect(panel_x, 0, RIGHT_W, H - band_h, fill=1, stroke=0)
        return band_h, panel_x

    def draw_creative_sidebar(panel_x, band_h):
        rx = panel_x + 14
        ry = H - band_h - 28

        if photo_img:
            draw_circular_photo(c, photo_img, panel_x + RIGHT_W / 2, ry + 5, 32)
            ry -= 50

        section_heading(c, rx, ry, "SKILLS", P1, 35)
        ry -= 20
        ry, _ = draw_skill_pills(c, cv.get("skills", []), rx, ry, RIGHT_W - 28, HexColor("#EDE9FE"), P1)

        if cv.get("education") and ry > 110:
            section_heading(c, rx, ry, "EDUCATION", P1, 58)
            ry -= 20
            draw_education(c, cv.get("education", []), rx, ry, RIGHT_W - 28, TEXT_DARK)
            return True
        return False

    def draw_continuation_shell(page_no):
        footer_brand(c, cv.get("is_premium", True))
        band_h = 58
        c.setFillColor(P1)
        c.rect(0, H - band_h, W, band_h, fill=1, stroke=0)
        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(white)
        c.drawString(28, H - 35, cv.get("name") or "Professional CV")
        c.setFont("Helvetica", 7)
        c.setFillColor(HexColor("#D0CDFF"))
        c.drawRightString(W - 24, H - 35, f"Page {page_no}")
        panel_x = W - RIGHT_W
        c.setFillColor(PANEL_BG)
        c.rect(panel_x, 0, RIGHT_W, H - band_h, fill=1, stroke=0)
        return band_h, panel_x

    band_h, panel_x = draw_page1_shell()
    edu_in_sidebar = draw_creative_sidebar(panel_x, band_h)

    lx = 28
    y = H - band_h - 28
    page_no = 1

    section_heading(c, lx, y, "SUMMARY", P1, 45)
    y -= 18
    y = draw_wrapped(c, cv.get("summary"), lx, y, LEFT_W, size=8.5, leading=13, colour=TEXT_MED) - 20

    section_heading(c, lx, y, "EXPERIENCE", P1, 60)
    y -= 18

    def _exp_newpage():
        nonlocal page_no
        c.showPage()
        page_no += 1
        bh, _px = draw_continuation_shell(page_no)
        return H - bh - 28
    for job in cv.get("experience", []):
        y = draw_role(c, job, lx, y, LEFT_W, P1, TEXT_DARK, TEXT_MED, "#7A7A8A", company_gap=9, new_page=_exp_newpage)

    # Extra senior-CV sections (render only if present)
    extra_sections = []
    if not edu_in_sidebar and cv.get("education"):
        edu_lines = [e.get("line") for e in cv.get("education", []) if e.get("line")]
        if edu_lines:
            extra_sections.append(("EDUCATION", edu_lines))
    extra_sections += [
        ("KEY ACHIEVEMENTS", cv.get("achievements", [])),
        ("PROFESSIONAL CERTIFICATIONS", cv.get("certifications", [])),
        ("SYSTEMS EXPERIENCE", cv.get("systems_experience", [])),
    ]
    for label, items in extra_sections:
        items = [clean_text(i) for i in (items or []) if clean_text(i)]
        if not items:
            continue
        if y < 104:
            c.showPage()
            page_no += 1
            band_h, panel_x = draw_continuation_shell(page_no)
            y = H - band_h - 28
        y -= 14
        section_heading(c, lx, y, label, P1, min(LEFT_W, 150))
        y -= 16
        for item in items:
            if y < 70:
                c.showPage()
                page_no += 1
                band_h, panel_x = draw_continuation_shell(page_no)
                y = H - band_h - 28
            y = draw_manual_bullet(c, item, lx, y, LEFT_W, colour=TEXT_MED, size=7.5, leading=9, bullet=True)
            y -= 0.8
        y -= 6

    if cv.get("references"):
        y -= 14
        if y >= 60:
            section_heading(c, lx, y, "REFERENCES", P1, 75)
            y -= 16
            draw_wrapped(c, cv.get("references"), lx, y, LEFT_W, size=8, leading=10, colour=TEXT_MED)

    c.save()
    buf.seek(0)
    return buf


def generate_impact_pdf(cv, colours):
    cv = normalise_cv(cv)
    HEADER_BG = safe_hex(colours.get("primary"), "#111827")
    TEAL = safe_hex(colours.get("accent"), "#0D9488")
    TEXT_DARK = "#111827"
    TEXT_MED = "#4B5563"
    RIGHT_W = 180

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    photo_img = decode_photo(cv.get("photo"))

    header_h = 100
    right_x = W - RIGHT_W - 4
    left_x = 28
    left_w = W - RIGHT_W - 52
    body_top = H - header_h - 22
    SIDEBAR_FLOOR = 40
    pending_skills = []

    def draw_page1_shell():
        footer_brand(c, cv.get("is_premium", True))

        c.setFillColor(HEADER_BG)
        c.rect(0, H - header_h, W, header_h, fill=1, stroke=0)

        c.setFont("Helvetica-Bold", 30)
        c.setFillColor(white)
        c.drawString(28, H - 48, cv.get("name") or "Professional CV")

        c.setFont("Helvetica", 7.5)
        c.setFillColor(HexColor("#9CA3AF"))
        c.drawString(28, H - 85, clean_join([cv.get("email"), cv.get("phone"), cv.get("location")], " | "))

        if photo_img:
            draw_circular_photo(c, photo_img, W - 65, H - header_h / 2, 31)

        c.setFillColor(HexColor("#F9FAFB"))
        c.rect(right_x - 10, 0, RIGHT_W + 14, body_top + 22, fill=1, stroke=0)

    def draw_impact_sidebar():
        rx = right_x + 6
        ry = body_top

        section_heading(c, rx, ry, "SKILLS", TEAL, 35)
        ry -= 22
        ry, remaining = draw_skill_pills(
            c, cv.get("skills", []), rx, ry, RIGHT_W - 24, HEADER_BG, white, min_y=SIDEBAR_FLOOR
        )
        pending_skills[:] = remaining

        # Only render education in the sidebar if every skill fit AND there is
        # genuine room. Otherwise education goes in the main column (never lost)
        # and the overflow skills continue in the page-2 sidebar.
        if not pending_skills and cv.get("education") and ry > 110:
            section_heading(c, rx, ry, "EDUCATION", TEAL, 58)
            ry -= 22
            draw_education(c, cv.get("education", []), rx, ry, RIGHT_W - 24, TEXT_DARK)
            return True
        return False

    def draw_continuation_shell(page_no):
        footer_brand(c, cv.get("is_premium", True))
        band_h = 58
        c.setFillColor(HEADER_BG)
        c.rect(0, H - band_h, W, band_h, fill=1, stroke=0)
        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(white)
        c.drawString(28, H - 35, cv.get("name") or "Professional CV")
        c.setFont("Helvetica", 7)
        c.setFillColor(HexColor("#9CA3AF"))
        c.drawRightString(W - 24, H - 35, f"Page {page_no}")
        c.setFillColor(HexColor("#F9FAFB"))
        c.rect(right_x - 10, 0, RIGHT_W + 14, H - band_h, fill=1, stroke=0)

        # Continue any skills that didn't fit on the previous page's sidebar so
        # the page-2+ sidebar is used instead of overflowing off page 1.
        if pending_skills:
            rx = right_x + 6
            ry = H - band_h - 18
            _, remaining = draw_skill_pills(
                c, pending_skills, rx, ry, RIGHT_W - 24, HEADER_BG, white, min_y=SIDEBAR_FLOOR
            )
            pending_skills[:] = remaining

        return band_h

    draw_page1_shell()
    edu_in_sidebar = draw_impact_sidebar()

    y = body_top
    page_no = 1
    section_heading(c, left_x, y, "SUMMARY", TEAL, 45)
    y -= 18
    y = draw_wrapped(c, cv.get("summary"), left_x, y, left_w, size=8.5, leading=13, colour=TEXT_MED) - 20

    section_heading(c, left_x, y, "EXPERIENCE", TEAL, 70)
    y -= 18

    def _exp_newpage():
        nonlocal page_no
        c.showPage()
        page_no += 1
        bh = draw_continuation_shell(page_no)
        return H - bh - 28
    for job in cv.get("experience", []):
        y = draw_role(
            c,
            job,
            left_x + 18,
            y,
            left_w - 18,
            TEAL,
            TEXT_DARK,
            TEXT_MED,
            "#9CA3AF",
            company_gap=11,
            new_page=_exp_newpage,
        )

    # Extra senior-CV sections (render only if present)
    extra_sections = []
    if not edu_in_sidebar and cv.get("education"):
        edu_lines = [e.get("line") for e in cv.get("education", []) if e.get("line")]
        if edu_lines:
            extra_sections.append(("EDUCATION", edu_lines))
    extra_sections += [
        ("KEY ACHIEVEMENTS", cv.get("achievements", [])),
        ("PROFESSIONAL CERTIFICATIONS", cv.get("certifications", [])),
        ("SYSTEMS EXPERIENCE", cv.get("systems_experience", [])),
    ]
    for label, items in extra_sections:
        items = [clean_text(i) for i in (items or []) if clean_text(i)]
        if not items:
            continue
        if y < 104:
            c.showPage()
            page_no += 1
            band_h = draw_continuation_shell(page_no)
            y = H - band_h - 28
        y -= 14
        section_heading(c, left_x + 18, y, label, TEAL, min(left_w - 18, 170))
        y -= 16
        for item in items:
            if y < 70:
                c.showPage()
                page_no += 1
                band_h = draw_continuation_shell(page_no)
                y = H - band_h - 28
            y = draw_manual_bullet(c, item, left_x + 18, y, left_w - 18, colour=TEXT_MED, size=7.5, leading=9, bullet=True)
            y -= 0.8
        y -= 6

    if cv.get("references"):
        y -= 14
        if y >= 60:
            section_heading(c, left_x + 18, y, "REFERENCES", TEAL, 75)
            y -= 16
            draw_wrapped(c, cv.get("references"), left_x + 18, y, left_w - 18, size=8, leading=10, colour=TEXT_MED)

    # Flush any sidebar skills still pending onto fresh continuation pages.
    # Covers the rare case of a very long skills list paired with short
    # experience that never triggered an experience-driven page break.
    while pending_skills:
        c.showPage()
        page_no += 1
        draw_continuation_shell(page_no)

    c.save()
    buf.seek(0)
    return buf


def generate_docx(cv):
    cv = normalise_cv(cv)
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run(cv.get("name") or "Professional CV")
    run.font.size = Pt(24)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    crun = contact.add_run(clean_join([cv.get("email"), cv.get("phone"), cv.get("location")], " | "))
    crun.font.size = Pt(9)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("SUMMARY", level=2)
    doc.add_paragraph(cv.get("summary", ""))

    doc.add_heading("EXPERIENCE", level=2)
    for job in cv.get("experience", []):
        jp = doc.add_paragraph()
        jr = jp.add_run(clean_join([job.get("title"), job.get("company")], " - "))
        jr.bold = True
        if job.get("dates"):
            dp = doc.add_paragraph(job.get("dates"))
            if dp.runs:
                dp.runs[0].font.size = Pt(9)
        for bullet in job.get("bullets", []):
            doc.add_paragraph(clean_text(bullet), style="List Bullet")

    doc.add_heading("SKILLS", level=2)
    doc.add_paragraph(cv.get("skills_csv", ""))

    doc.add_heading("EDUCATION", level=2)
    for edu in cv.get("education", []):
        if edu.get("line"):
            doc.add_paragraph(edu.get("line"))

    if cv.get("certifications"):
        doc.add_heading("CERTIFICATIONS", level=2)
        for cert in cv.get("certifications", []):
            doc.add_paragraph(cert)

    if cv.get("achievements"):
        doc.add_heading("KEY ACHIEVEMENTS", level=2)
        for ach in cv.get("achievements", []):
            doc.add_paragraph(clean_text(ach), style="List Bullet")

    if cv.get("systems_experience"):
        doc.add_heading("SYSTEMS EXPERIENCE", level=2)
        for sys_item in cv.get("systems_experience", []):
            doc.add_paragraph(clean_text(sys_item), style="List Bullet")

    if cv.get("references"):
        doc.add_heading("REFERENCES", level=2)
        doc.add_paragraph(clean_text(cv.get("references")))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf




def no_store_response(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    response.headers["X-RoleAlign-API-Version"] = API_VERSION
    return response


def extract_cv_payload(data):
    raw_cv = data.get("cv_data")

    if raw_cv is None:
        raw_cv = data.get("cv")

    if not isinstance(raw_cv, dict) or not raw_cv:
        raise ValueError("Missing cv_data payload")

    has_text_content = any(
        clean_text(raw_cv.get(key))
        for key in ["name", "full_name", "summary", "email", "phone", "location"]
    )
    has_structured_content = bool(raw_cv.get("experience")) or bool(raw_cv.get("skills")) or bool(raw_cv.get("education"))

    if not has_text_content and not has_structured_content:
        raise ValueError("Invalid cv_data payload")

    return raw_cv

@app.route("/generate-pdf", methods=["POST"])
def gen_pdf():
    try:
        data = request.get_json(silent=True) or {}
        raw_cv = extract_cv_payload(data)
        template = (data.get("template") or data.get("template_id") or "executive").lower()
        cv = raw_cv
        colours = data.get("colours") or data.get("colors") or {}
        render_version = data.get("render_version") or cv.get("render_version") or API_VERSION
        cv = normalise_cv({**cv, "render_version": render_version})

        print(
            {
                "event": "generate_pdf",
                "api_version": API_VERSION,
                "template": template,
                "render_version": render_version,
                "has_photo": bool(cv.get("photo")),
                "skills_count": len(cv.get("skills", [])),
                "experience_count": len(cv.get("experience", [])),
                "has_skills_csv": bool(cv.get("skills_csv")),
                "has_experience_header": any(j.get("header") for j in cv.get("experience", [])),
            }
        )

        if template == "executive":
            buf = generate_executive_pdf(cv, colours)
        elif template == "creative":
            buf = generate_creative_pdf(cv, colours)
        elif template == "impact":
            buf = generate_impact_pdf(cv, colours)
        elif template == "starter":
            cv["is_premium"] = False
            buf = generate_starter_pdf(cv, colours)
        else:
            return jsonify({"error": f"Invalid template: {template}"}), 400

        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", cv.get("name") or "CV").strip("_") or "CV"
        response = send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"CV_{template}_{safe_name}.pdf",
        )
        return no_store_response(response)
    except ValueError as e:
        print(f"[generate_pdf_validation_error] {e}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        print(f"[generate_pdf_error] {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/generate-docx", methods=["POST"])
def gen_docx():
    try:
        data = request.get_json(silent=True) or {}
        raw_cv = extract_cv_payload(data)
        cv = normalise_cv(raw_cv)

        print(
            {
                "event": "generate_docx",
                "api_version": API_VERSION,
                "render_version": data.get("render_version") or cv.get("render_version") or API_VERSION,
                "skills_count": len(cv.get("skills", [])),
                "experience_count": len(cv.get("experience", [])),
                "has_photo": bool(cv.get("photo")),
            }
        )

        buf = generate_docx(cv)
        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", cv.get("name") or "CV").strip("_") or "CV"
        response = send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"CV_{safe_name}.docx",
        )
        return no_store_response(response)
    except ValueError as e:
        print(f"[generate_docx_validation_error] {e}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        print(f"[generate_docx_error] {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
